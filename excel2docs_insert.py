from pathlib import Path
from time import time
from docx import Document
import polars as pl
from docx.document import Document as DocumentObject
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Pt, RGBColor
from docx.table import _Cell
import openpyxl
from humanize import intcomma

SAVE_PATH = Path("output")
SAVE_PATH.mkdir(exist_ok=True)
ORIGIN_PATH = Path('шаблоны', 'task3')
OUTPUT_DIR = Path("output")
OUTPUT_DIR.mkdir(exist_ok=True)
timestamp = str(int(time()))
tmp_dir = OUTPUT_DIR / timestamp
tmp_dir.mkdir()

def has_red_marker(cell: _Cell, marker: str):
    for p in cell.paragraphs:
        if marker in p.text:
            colors = set([p.runs[c].font.color.rgb for c in range(len(p.runs))])
            if len(colors) == 1 and colors.pop() == RGBColor(0xee, 0x00, 0x00):
                return True
    return None
def search_red_marker(cell: _Cell):
    for p in cell.paragraphs:
        # print('p.text = ', p.text)
        colors = set([p.runs[c].font.color.rgb for c in range(len(p.runs))])
        if len(colors) == 1 and colors.pop() == RGBColor(0xee, 0x00, 0x00):
            return p.text.strip()
    return None

def insert_table(doc: DocumentObject, df: pl.DataFrame, marker: str):
    all_cells = (
        cell
        for tbl in doc.tables
        for row in tbl.rows
        for cell in row.cells
    )

    # находим первую ячейку с маркером L6
    try:
        cell = next(cell for cell in all_cells if marker in cell.text)
    except StopIteration:
        return None
    
    if has_red_marker(cell, marker):
        # очистить маркер
        cell.text = cell.text.replace(marker, "")
        # вставить вложенную таблицу
        nested = cell.add_table(rows=1, cols=df.width)
        nested.style = "Table Grid"
        nested.alignment = WD_TABLE_ALIGNMENT.LEFT
        # nested.autofit = False

        # шапка

        hdr = nested.rows[0].cells
        for i, col in enumerate(df.columns):
            hdr[i].text = str(col).strip()
        # данные
        for data_row in df.iter_rows():
            new_cells = nested.add_row().cells
            for i, val in enumerate(data_row):
                new_cells[i].text = intcomma(str(val).strip()).replace(",", " ").replace(".", ",") if i >= 2 else str(val).strip()
        for hdr_cell in nested.rows[0].cells:
            for p in hdr_cell.paragraphs:
                pf = p.paragraph_format
                pf.first_line_indent = Pt(0)
                pf.left_indent = Pt(0)

        # и то же для всех строк с данными
        for data_row in nested.rows[1:]:
            for cell in data_row.cells:
                for p in cell.paragraphs:
                    pf = p.paragraph_format
                    pf.first_line_indent = Pt(0)
                    pf.left_indent = Pt(0)

        return doc
    raise RuntimeError(f"Маркер {marker} не найден ни в одной ячейке таблиц.")

if __name__ == "__main__":
    doc_path = ORIGIN_PATH / 'шаблон.docx'
    excel_path = ORIGIN_PATH / 'РД Выборка Индо Банк 24-09.xlsx'
    doc = Document(doc_path)
    wb = openpyxl.load_workbook(excel_path, data_only=True, read_only=True)
    ws = wb['Приложение_ОСВ']

    # 2. Собираем номера строк, где хотя бы одна ячейка залита жёлтым (#FFFF00 или ColorIndex 6)
    yellow_rows = {"лист": [], "Лицевой счет": [], "Наименование счета": []}
    for row in ws.iter_rows(min_row=2):  # header_row=1 → данные с физической строки 2
        for i, cell in enumerate(row):
            fg = cell.fill.fgColor
            # openpyxl хранит цвет в разных форматах, но в большинстве случаев .rgb == 'FFFFFF00' или 'FF0000FF'
            rgb = getattr(fg, 'rgb', None)
            # некоторые файлы могут использовать indexed цвет
            idx = getattr(fg, 'index', None)
            if rgb and rgb.upper().endswith('FFFF00'):
                yellow_rows["лист"].append(str(row[i - 1].value))
                yellow_rows["Лицевой счет"].append(str(row[i].value))
                yellow_rows["Наименование счета"].append(str(row[i + 1].value))
                break  # эту строку уже отметили, идём дальше
    # print('yellow_rows = ', yellow_rows)
    df = pl.DataFrame(yellow_rows)
    # print(df)

    all_cells = (
        cell
        for tbl in doc.tables
        for row in tbl.rows
        for cell in row.cells
    )

    for cell in all_cells:
        marker = search_red_marker(cell)
        if marker:
            temp_df = df.filter(pl.nth(0) == marker).drop(pl.nth(0))
            print(marker, temp_df)
            doc = insert_table(doc, temp_df, marker)
    doc.save(tmp_dir / f'выход_{doc_path.name}')
