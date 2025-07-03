import tomllib
from pathlib import Path
import polars as pl
from docx import Document
from docx.document import Document as DocumentObject
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from pydantic import BaseModel
from time import time
import polars.selectors as sc
from humanize import intcomma
from docx.shared import Pt
from docx.table import _Cell
import openpyxl
from copy import copy
from openpyxl.worksheet.worksheet import Worksheet
from openpyxl.utils import column_index_from_string

OUTPUT_DIR = Path("output")
START_PART = "_ВСТАВКА_"

class OfficeConfig(BaseModel):
    filename: Path
    sheet_names: list[str] | None = None


class Processor:
    word: OfficeConfig
    excel: OfficeConfig

    def __init__(self) -> None:
        OUTPUT_DIR.mkdir(exist_ok=True)

        with open("config.toml", "rb") as f:
            cfg = tomllib.load(f)
        self.word = OfficeConfig.model_validate(cfg["word"])
        self.excel = OfficeConfig.model_validate(cfg["excel"])
        pass
    
    # TASK2
    def make_word(self, word_filename, excel_filename) -> str:
        timestamp = str(int(time()))
        tmp_dir = OUTPUT_DIR / timestamp
        tmp_dir.mkdir()

        # 1. конфиг
        print("filenames = ", word_filename, excel_filename)

        self.word.filename = Path(str(word_filename))
        self.excel.filename = Path(str(excel_filename))
        df = pl.read_excel(self.excel.filename, sheet_name=self.excel.sheet_names[0])
        df = df.with_columns(sc.by_dtype(pl.Float64).round(2))

        # 3. Открываем Word
        doc = Document(self.word.filename)
        tmp = tmp_dir / f"temp_{self.word.filename.name}"

        # 4. Ищем L6 внутри ячеек всех таблиц
        doc = insert_l6_table(doc, df)
        if not doc:
            raise RuntimeError("Маркер L6 не найден ни в одной ячейке таблиц.")
        doc.save(tmp)

        df = pl.read_excel(self.excel.filename, sheet_name=self.excel.sheet_names[1])
        doc = insert_k_table(doc, df)
        if not doc:
            raise RuntimeError("Маркер L6 не найден ни в одной ячейке таблиц.")
        doc.save(tmp)

        df = pl.read_excel(self.excel.filename, sheet_name=self.excel.sheet_names[2])
        doc = insert_d_table(doc, df)
        if not doc:
            raise RuntimeError(
                "Маркер Общая сумма СПОД не найден ни в одной ячейке таблиц."
            )
        # 5. Сохраняем документ
        doc.save(tmp)
        tmp = tmp.relative_to(".")
        print("tmp = ", tmp, str(tmp))
        return str(tmp)
    # TASK 3
    def excel2word_insert(self, word_filename: str, excel_filename: str):
        timestamp = str(int(time()))
        tmp_dir = OUTPUT_DIR / timestamp
        tmp_dir.mkdir()

        word_filename = Path(str(word_filename))
        excel_filename = Path(str(excel_filename))
        doc = Document(word_filename)
        wb = openpyxl.load_workbook(excel_filename, data_only=True, read_only=True)
        ws = wb["Приложение_ОСВ"]

        # 2. Собираем номера строк, где хотя бы одна ячейка залита жёлтым (#FFFF00 или ColorIndex 6)
        yellow_rows = {"лист": [], "Лицевой счет": [], "Наименование счета": []}
        for row in ws.iter_rows(
            min_row=2
        ):  # header_row=1 → данные с физической строки 2
            for i, cell in enumerate(row):
                # print(i, 'value: ', cell.value) if i % 5 == 0 else None
                if not hasattr(cell.fill, 'fgColor'):
                    continue
                fg = cell.fill.fgColor
                # openpyxl хранит цвет в разных форматах, но в большинстве случаев .rgb == 'FFFFFF00' или 'FF0000FF'
                rgb = getattr(fg, "rgb", None)
                # некоторые файлы могут использовать indexed цвет
                if rgb and rgb.upper().endswith("FFFF00"):
                    yellow_rows["лист"].append(str(row[i - 1].value))
                    yellow_rows["Лицевой счет"].append(str(row[i].value))
                    yellow_rows["Наименование счета"].append(str(row[i + 1].value))
                    break  # эту строку уже отметили, идём дальше
        # print('yellow_rows = ', yellow_rows)
        df = pl.DataFrame(yellow_rows)
        # print(df)

        # Перебираем таблицы и строки, чтобы иметь доступ к row
        for tbl in doc.tables:
            # Копируем список строк, чтобы безопасно удалять во время итерации
            for row in list(tbl.rows):
                found_marker = False
                for cell in row.cells:
                    marker = search_text_marker(cell)
                    if marker:
                        temp_df = df.filter(pl.nth(0) == marker).drop(pl.nth(0))
                        if temp_df.is_empty():
                            # Удаляем строку, если DataFrame пуст
                            tbl._tbl.remove(row._tr)
                            found_marker = True
                            break
                        print(marker, temp_df)
                        temp_doc = insert_table(doc, temp_df, marker)
                        if temp_doc:
                            doc = temp_doc
                        else:
                            tbl._tbl.remove(row._tr)
                        found_marker = True
                        break
                if not found_marker:
                    # Если не найден ни один marker в строке, удаляем строку
                    continue
        tmp = tmp_dir / f"выход_{word_filename.name}"
        doc.save(tmp)
        tmp = tmp.relative_to(".")
        return str(tmp)

    def copy_ws(self, origin_filename, target_filename):
        from openpyxl import load_workbook

        origin_filename = Path(str(origin_filename))
        target_filename = Path(str(target_filename))
        timestamp = str(int(time()))
        tmp_dir = OUTPUT_DIR / timestamp
        tmp_dir.mkdir()
        origin_wb = load_workbook(origin_filename)
        target_wb = load_workbook(target_filename)

        for sheet_name in reversed(origin_wb.sheetnames):
            # Удаляем существующий лист, если он есть
            if sheet_name in target_wb.sheetnames:
                target_wb.remove(target_wb[sheet_name])
            # Создаём новый лист в позиции 0 (в начале)
            dst = target_wb.create_sheet(title=sheet_name, index=0)
            src = origin_wb[sheet_name]
            copy_sheet(src, dst)

        # Optionally remove the default 'Sheet' if it's empty and not in origin
        if "Sheet" in target_wb.sheetnames and "Sheet" not in origin_wb.sheetnames:
            target_wb.remove(target_wb["Sheet"])

        # Save changes
        timestamp = str(int(time()))
        tmp = tmp_dir / timestamp
        tmp.mkdir()
        tmp = tmp / f"target_{target_filename.name}"

        target_wb.save(tmp)
        return str(tmp.relative_to("."))


def insert_l6_table(doc: DocumentObject, df: pl.DataFrame):
    all_cells = (cell for tbl in doc.tables for row in tbl.rows for cell in row.cells)

    # находим первую ячейку с маркером L6
    try:
        cell = next(cell for cell in all_cells if "L6" in cell.text)
    except StopIteration:
        return None
    clear_cell_shading(cell)
    if "L6" in cell.text:
        # очистить маркер
        cell.text = ""
        # вставить вложенную таблицу
        nested = cell.add_table(rows=1, cols=len(df.columns))
        nested.style = "Table Grid"
        nested.alignment = WD_TABLE_ALIGNMENT.LEFT
        # nested.autofit = False

        # шапка

        hdr = nested.rows[0].cells
        for i, col in enumerate(df.columns):
            hdr[i].text = str(col).strip()
        # данные
        for data_row in df.iter_rows():
            new_cells = nested.add_row().cells
            for i, val in enumerate(data_row):
                new_cells[i].text = (
                    intcomma(str(val).strip()).replace(",", " ").replace(".", ",")
                    if i >= 2
                    else str(val).strip()
                )
        for hdr_cell in nested.rows[0].cells:
            for p in hdr_cell.paragraphs:
                pf = p.paragraph_format
                pf.first_line_indent = Pt(0)
                pf.left_indent = Pt(0)

        # и то же для всех строк с данными
        for data_row in nested.rows[1:]:
            for cell in data_row.cells:
                for p in cell.paragraphs:
                    pf = p.paragraph_format
                    pf.first_line_indent = Pt(0)
                    pf.left_indent = Pt(0)

        return doc


def insert_k_table(doc: DocumentObject, df: pl.DataFrame):
    df = df.with_columns(sc.by_dtype(pl.Float64).round(4))
    # 1) Находим таблицу и номер заголовочной строки
    target_table = None
    header_row_idx = None

    for tbl in doc.tables:
        for idx, row in enumerate(tbl.rows):
            # проверяем, есть ли в этой строке нужная ячейка
            if any("Общая сумма СПОД" in cell.text for cell in row.cells):
                target_table = tbl
                header_row_idx = idx
                break
        if target_table:
            break

    if not target_table:
        # таблица не найдена — выходим
        return

    # 2) Берём значения из df
    # Предположим, что у вас в df есть колонка "value", и в ней ровно столько строк,
    # сколько строк таблицы (начиная с header_row_idx), куда нужно записать.
    values = df.row(0)

    # 3) Проверяем, хватает ли строк в таблице.
    # Нужны строки от header_row_idx до header_row_idx + len(values) - 1
    needed = header_row_idx + len(values) - len(target_table.rows)
    for _ in range(needed):
        target_table.add_row()

    # 4) Записываем по одной строке в ячейку (row_idx, col_idx=1)
    # print(target_table.rows)
    # print(len(target_table.rows))

    row_indices = []
    for row_idx in range(header_row_idx, header_row_idx + len(target_table.rows) - 1):
        # проверяем первый столбец на непустоту
        # print('row_idx = ', row_idx)
        if (
            not target_table.cell(row_idx, 0).text.strip()
            or not target_table.cell(row_idx, 1).text.strip()
        ):
            continue
        # проверяем второй столбец на "Д"
        if target_table.cell(row_idx, 1).text.startswith("Д"):
            continue
        row_indices.append(row_idx)

    # 2) Теперь просто zip по values и row_indices
    for val, row_idx in zip(values, row_indices):
        cell = target_table.cell(row_idx, 1)
        cell.text = intcomma(str(val)).replace(",", " ").replace(".", ",")
    return doc


def insert_d_table(doc: DocumentObject, df: pl.DataFrame):
    df = df.with_columns(sc.by_dtype(pl.Float64).round(2))
    # 1) Находим таблицу и номер заголовочной строки
    target_table = None
    header_row_idx = None

    for tbl in doc.tables:
        for idx, row in enumerate(tbl.rows):
            # проверяем, есть ли в этой строке нужная ячейка
            if any("Общая сумма СПОД" in cell.text for cell in row.cells):
                target_table = tbl
                header_row_idx = idx
                break
        if target_table:
            break

    if not target_table:
        # таблица не найдена — выходим
        return

    # 2) Берём значения из df
    values = df.to_series(1).to_list()

    # 3) Проверяем, хватает ли строк в таблице.
    # Нужны строки от header_row_idx до header_row_idx + len(values) - 1
    needed = header_row_idx + len(values) - len(target_table.rows)
    for _ in range(needed):
        target_table.add_row()

    # 4) Записываем по одной строке в ячейку (row_idx, col_idx=1)
    # print(target_table.rows)
    # print(len(target_table.rows))

    row_indices = []
    for row_idx in range(header_row_idx, header_row_idx + len(target_table.rows) - 1):
        # проверяем первый столбец на непустоту
        # print('row_idx = ', row_idx)
        if (
            not target_table.cell(row_idx, 0).text.strip()
            or not target_table.cell(row_idx, 1).text.strip()
        ):
            continue
        # проверяем второй столбец на "Д"
        if not target_table.cell(row_idx, 1).text.startswith("Д"):
            continue
        row_indices.append(row_idx)

    # 2) Теперь просто zip по values и row_indices
    for val, row_idx in zip(values, row_indices):
        cell = target_table.cell(row_idx, 1)
        cell.text = intcomma(str(val)).replace(",", " ").replace(".", ",")
    return doc


def clear_cell_shading(cell):
    """
    Убирает любую заливку из ячейки cell:
    удаляет все <w:shd> в свойствах ячейки.
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for shd in tcPr.findall(qn("w:shd")):
        tcPr.remove(shd)


def has_text_marker(cell: _Cell, marker: str, start_par: str = START_PART):
    for p in cell.paragraphs:
        if marker in p.text:
            if p.text.startswith(start_par):
                return True
    return None


def search_text_marker(cell: _Cell, start_par: str = START_PART):
    for p in cell.paragraphs:
        # print('p.text = ', p.text)
        if p.text.startswith(start_par):
            return p.text.strip().removeprefix(start_par)
    return None


def insert_table(doc: DocumentObject, df: pl.DataFrame, marker: str):
    all_cells = (cell for tbl in doc.tables for row in tbl.rows for cell in row.cells)

    # находим первую ячейку с маркером L6
    try:
        cell = next(cell for cell in all_cells if marker in cell.text)
    except StopIteration:
        return None

    if has_text_marker(cell, START_PART + marker):
        # очистить маркер
        cell.text = cell.text.replace(START_PART + marker, "")
        # вставить вложенную таблицу
        nested = cell.add_table(rows=1, cols=df.width)
        nested.style = "Table Grid"
        nested.alignment = WD_TABLE_ALIGNMENT.LEFT
        # nested.autofit = False

        # шапка

        hdr = nested.rows[0].cells
        for i, col in enumerate(df.columns):
            hdr[i].text = str(col).strip()
        # данные
        for data_row in df.iter_rows():
            new_cells = nested.add_row().cells
            for i, val in enumerate(data_row):
                new_cells[i].text = (
                    intcomma(str(val).strip()).replace(",", " ").replace(".", ",")
                    if i >= 2
                    else str(val).strip()
                )
        for hdr_cell in nested.rows[0].cells:
            for p in hdr_cell.paragraphs:
                pf = p.paragraph_format
                pf.first_line_indent = Pt(0)
                pf.left_indent = Pt(0)

        # и то же для всех строк с данными
        for data_row in nested.rows[1:]:
            for cell in data_row.cells:
                for p in cell.paragraphs:
                    pf = p.paragraph_format
                    pf.first_line_indent = Pt(0)
                    pf.left_indent = Pt(0)

        return doc
    print(f"Маркер {marker} не найден ни в одной ячейке таблиц.")


def copy_sheet(src_ws: Worksheet, dst_ws: Worksheet):
    # Копируем значения и стили
    for row in src_ws.iter_rows():
        for cell in row:
            row_idx = cell.row
            # для обычных ячеек берем col_idx, для merged — преобразуем letter -> index
            if hasattr(cell, "col_idx"):
                col_idx = cell.col_idx
            else:
                col = cell.column  # Может быть буквой
                col_idx = col if isinstance(col, int) else column_index_from_string(col)
            new_cell = dst_ws.cell(row=row_idx, column=col_idx, value=cell.value)

            if cell.has_style:
                new_cell.font = copy(cell.font)
                new_cell.border = copy(cell.border)
                new_cell.fill = copy(cell.fill)
                # Прямо присваиваем строковый код формата
                new_cell.number_format = cell.number_format
                new_cell.protection = copy(cell.protection)
                new_cell.alignment = copy(cell.alignment)
    # Воссоздаём merged ranges
    for merged_range in src_ws.merged_cells.ranges:
        dst_ws.merge_cells(str(merged_range))
    # Копируем размеры строк
    for idx, dim in src_ws.row_dimensions.items():
        dst_ws.row_dimensions[idx].height = dim.height
    # Копируем ширины столбцов
    for idx, dim in src_ws.column_dimensions.items():
        dst_ws.column_dimensions[idx].width = dim.width
    # Копируем остальные свойства листа
    dst_ws.sheet_format = src_ws.sheet_format
    dst_ws.sheet_properties = src_ws.sheet_properties
    dst_ws.page_margins = src_ws.page_margins
    dst_ws.page_setup = src_ws.page_setup
    dst_ws.print_options = src_ws.print_options

    # --- Новый блок: копируем условное форматирование ---
    # Простой (хлороформатный) способ — скопировать «правила» целиком:
    dst_ws.conditional_formatting._cf_rules = copy(
        src_ws.conditional_formatting._cf_rules
    )
