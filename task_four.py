from docx.text.paragraph import Paragraph
from docx.table import Table
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.shared import Pt
import polars as pl
from humanize import intcomma
from pathlib import Path
from time import time

START_PART = "_ВСТАВКА_"
OUTPUT_DIR = Path("output")
OUTPUT_DIR.mkdir(exist_ok=True)

def insert_table_after(paragraph: Paragraph, df: pl.DataFrame, marker: str = START_PART) -> Table:
    """
    Вставляет таблицу сразу после параграфа с маркером.
    """
    # 1. Убираем маркер из текста параграфа
    if marker in paragraph.text:
        paragraph.text = paragraph.text.replace(marker, "").strip()

    # 2. Создаем новую таблицу
    rows = df.height + 1  # +1 на заголовок
    cols = df.width
    # Вычисляем доступную ширину страницы и используем ее для создания таблицы
    section = paragraph.part.document.sections[0]
    available_width = section.page_width - section.left_margin - section.right_margin
    table = paragraph._parent.add_table(rows=rows, cols=cols, width=available_width)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # 3. Перемещаем таблицу в нужное место (после параграфа)
    p = paragraph._element
    parent = p.getparent()
    parent.insert(parent.index(p) + 1, table._element)

    # 4. Заполняем таблицу данными
    # Шапка
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(df.columns):
        hdr_cells[i].text = str(col).strip()

    # Данные - ИСПРАВЛЕНА ЛОГИКА ЗАПОЛНЕНИЯ
    for i, data_row in enumerate(df.iter_rows()):
        row_cells = table.rows[i + 1].cells  # Используем существующие строки, а не добавляем новые
        for j, val in enumerate(data_row):
            text = str(val).strip() if val is not None else ""
            try:
                text = intcomma(str(val).strip()).replace(",", " ").replace(".", ",") if j - len(data_row) >= -4 else str(val).strip()
            except (ValueError, TypeError):
                pass  # Оставляем как есть, если не число
            row_cells[j].text = text

    # 5. Убираем отступы
    for row in table.rows:
        for cell in row.cells:
            for p_cell in cell.paragraphs:
                p_cell.paragraph_format.first_line_indent = Pt(0)
                p_cell.paragraph_format.left_indent = Pt(0)

    return table


def insert_tables_with_filter(word_filename: str, excel_filename: str):
    timestamp = str(int(time()))
    tmp_dir = OUTPUT_DIR / timestamp
    tmp_dir.mkdir()

    word_filename: Path = Path(str(word_filename))
    excel_filename: Path = Path(str(excel_filename))  
    try:
        doc = Document(word_filename)
        # sheet_id=0 читает все листы
        dfs_dict = pl.read_excel(excel_filename, sheet_id=0, infer_schema_length=0,
                                 read_options={"header_row": 2})
    except FileNotFoundError as e:
        print(f"Ошибка: Файл не найден - {e}")
        return

    sheet_names = list(dfs_dict.keys())
    print(f"Найденные листы в Excel: {sheet_names}")

    # Сначала найдем все параграфы-плейсхолдеры, чтобы избежать проблем при итерации
    placeholder_paragraphs = [p for p in doc.paragraphs if p.text.strip().startswith(START_PART)]

    print(f"Найдено {len(placeholder_paragraphs)} меток для вставки таблиц в Word.")

    if not placeholder_paragraphs:
        print(f"В документе не найдено меток для вставки таблиц. Проверьте текст на наличие '{START_PART}'.")
        return

    if not sheet_names:
        print("В Excel файле не найдено листов для вставки.")
        return

    # Вставляем таблицы
    tables_to_insert_count = min(len(placeholder_paragraphs), len(sheet_names))
    if len(placeholder_paragraphs) != len(sheet_names):
        print(f"Внимание: Количество меток ({len(placeholder_paragraphs)}) не совпадает с количеством листов ({len(sheet_names)}). Будет вставлено {tables_to_insert_count} таблиц.")

    for i in range(tables_to_insert_count):
        p = placeholder_paragraphs[i]
        sheet_name = sheet_names[i]
        df = dfs_dict.get(sheet_name).filter(~pl.any_horizontal(pl.nth(-1, -2, -3) == '0'))

        print(f"Вставка таблицы с листа '{sheet_name}'...")
        insert_table_after(p, df)

    tmp = tmp_dir / f"выход_{word_filename.name}"
    tmp = tmp.relative_to(".")

    doc.save(tmp)
    return str(tmp)

